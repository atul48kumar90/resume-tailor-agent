from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from agents.resume_formatter import format_resume_text


def export_pdf(resume_text: str) -> BytesIO:
    """
    Generates a clean, ATS-safe PDF resume.
    Note: This function takes resume_text (string), not resume dict.
    For resume dict, use exporters/pdf_exporter.py
    """
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36,
    )

    styles = getSampleStyleSheet()
    normal = styles["Normal"]

    story = []

    for line in resume_text.split("\n"):
        if not line.strip():
            story.append(Spacer(1, 8))
        else:
            story.append(Paragraph(line, normal))

    doc.build(story)
    buffer.seek(0)
    return buffer


def export_docx(resume_text: str) -> BytesIO:
    """
    Exports resume text to DOCX format as BytesIO.
    Note: This function takes resume_text (string), not resume dict.
    For resume dict, use exporters/docx_exporter.py
    """
    from docx import Document
    
    doc = Document()
    
    for line in resume_text.split("\n"):
        if line.strip():
            if line.strip().upper() in ["SUMMARY", "SKILLS", "EXPERIENCE"]:
                doc.add_heading(line.strip(), level=1)
            elif line.strip().startswith("- "):
                doc.add_paragraph(line.strip()[2:], style="List Bullet")
            else:
                doc.add_paragraph(line.strip())
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
