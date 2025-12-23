# agents/ats_format_validator.py
"""
ATS format validation - checks if exported resumes are ATS-parseable.
"""
from typing import Dict, List, Any, IO
import re
import logging

logger = logging.getLogger(__name__)


def validate_ats_format(file_path: str = None, file_obj: IO = None, file_type: str = "pdf") -> Dict[str, Any]:
    """
    Validate if a resume file is ATS-friendly.
    
    Args:
        file_path: Path to file (if file_obj not provided)
        file_obj: File object (if file_path not provided)
        file_type: File type ('pdf' or 'docx')
    
    Returns:
        Validation results with issues and recommendations
    """
    issues = []
    warnings = []
    recommendations = []
    
    if file_type == "pdf":
        return _validate_pdf_ats(file_path, file_obj, issues, warnings, recommendations)
    elif file_type == "docx":
        return _validate_docx_ats(file_path, file_obj, issues, warnings, recommendations)
    else:
        return {
            "valid": False,
            "score": 0,
            "issues": ["Unsupported file type for ATS validation"],
            "warnings": [],
            "recommendations": []
        }


def _validate_pdf_ats(
    file_path: str,
    file_obj: IO,
    issues: List[str],
    warnings: List[str],
    recommendations: List[str]
) -> Dict[str, Any]:
    """Validate PDF for ATS compatibility."""
    try:
        import pdfplumber
        from PyPDF2 import PdfReader
        
        # Use file_obj if provided, otherwise open file_path
        if file_obj:
            # Reset file position
            file_obj.seek(0)
            pdf_file = file_obj
        else:
            pdf_file = open(file_path, 'rb')
        
        try:
            # Check if PDF is text-based (not scanned image)
            reader = PdfReader(pdf_file)
            pdf_file.seek(0)  # Reset for pdfplumber
            
            has_text = False
            total_pages = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and len(text.strip()) > 10:
                    has_text = True
                    break
            
            if not has_text:
                issues.append("PDF appears to be image-based (scanned). ATS systems cannot parse images.")
                recommendations.append("Convert scanned PDF to text-based PDF using OCR or recreate the document.")
            
            # Check for text extraction
            with pdfplumber.open(pdf_file) as pdf:
                extractable_text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extractable_text += page_text
                
                if len(extractable_text.strip()) < 50:
                    issues.append("Very little extractable text found. PDF may be image-based or corrupted.")
                
                # Check for common ATS-breaking elements
                _check_ats_breaking_elements(extractable_text, issues, warnings, recommendations)
            
            # Check PDF structure
            if total_pages > 2:
                warnings.append(f"Resume is {total_pages} pages. Most ATS systems prefer 1-2 pages.")
                recommendations.append("Consider condensing to 1-2 pages for better ATS compatibility.")
            
        finally:
            if not file_obj and pdf_file:
                pdf_file.close()
        
    except Exception as e:
        logger.error(f"PDF validation error: {e}", exc_info=True)
        issues.append(f"Error validating PDF: {str(e)}")
    
    return _format_validation_result(issues, warnings, recommendations, "pdf")


def _validate_docx_ats(
    file_path: str,
    file_obj: IO,
    issues: List[str],
    warnings: List[str],
    recommendations: List[str]
) -> Dict[str, Any]:
    """Validate DOCX for ATS compatibility."""
    try:
        from docx import Document
        
        # Use file_obj if provided, otherwise open file_path
        if file_obj:
            doc = Document(file_obj)
        else:
            doc = Document(file_path)
        
        # Extract all text
        full_text = "\n".join([para.text for para in doc.paragraphs])
        
        # Check for tables (some ATS systems struggle with complex tables)
        table_count = len(doc.tables)
        if table_count > 3:
            warnings.append(f"Document contains {table_count} tables. Complex tables may not parse correctly in ATS.")
            recommendations.append("Consider simplifying table structure or converting to plain text.")
        
        # Check for images
        # Note: python-docx doesn't directly expose images, but we can check for image references
        if len(full_text.strip()) < 50:
            issues.append("Very little text found. Document may be mostly images or empty.")
        
        # Check for text boxes (harder for ATS to parse)
        # This is harder to detect with python-docx, but we can check formatting
        
        # Check for ATS-breaking elements in text
        _check_ats_breaking_elements(full_text, issues, warnings, recommendations)
        
        # Check document length
        word_count = len(full_text.split())
        if word_count > 1000:
            warnings.append(f"Document is very long ({word_count} words). ATS systems prefer concise resumes.")
        elif word_count < 100:
            warnings.append("Document is very short. May be missing important information.")
        
    except Exception as e:
        logger.error(f"DOCX validation error: {e}", exc_info=True)
        issues.append(f"Error validating DOCX: {str(e)}")
    
    return _format_validation_result(issues, warnings, recommendations, "docx")


def _check_ats_breaking_elements(
    text: str,
    issues: List[str],
    warnings: List[str],
    recommendations: List[str]
):
    """Check text for ATS-breaking elements."""
    text_lower = text.lower()
    
    # Check for headers/footers (often problematic)
    if "page" in text_lower and re.search(r'page\s+\d+', text_lower):
        warnings.append("Page numbers detected. Some ATS systems may include them in parsed text.")
    
    # Check for special characters that might break parsing
    problematic_chars = ['•', '●', '▪', '▫', '→', '←']
    found_chars = [char for char in problematic_chars if char in text]
    if found_chars:
        warnings.append(f"Special characters detected: {', '.join(set(found_chars))}. Some ATS systems may not parse these correctly.")
        recommendations.append("Consider using standard bullet points (- or *) instead of special characters.")
    
    # Check for email/phone formatting
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    
    if not emails:
        warnings.append("No email address found. Contact information is important for ATS systems.")
    
    if not phones:
        warnings.append("No phone number found. Contact information is important for ATS systems.")
    
    # Check for keywords density (too many might be keyword stuffing)
    # This is a simple check - count repeated words
    words = text_lower.split()
    word_freq = {}
    for word in words:
        if len(word) > 4:  # Only check longer words
            word_freq[word] = word_freq.get(word, 0) + 1
    
    high_freq_words = [word for word, count in word_freq.items() if count > 20]
    if high_freq_words:
        warnings.append(f"Some words appear very frequently: {', '.join(high_freq_words[:3])}. May indicate keyword stuffing.")


def _format_validation_result(
    issues: List[str],
    warnings: List[str],
    recommendations: List[str],
    file_type: str
) -> Dict[str, Any]:
    """Format validation results."""
    # Calculate score (0-100)
    score = 100
    
    # Deduct points for issues
    score -= len(issues) * 20  # Each issue costs 20 points
    
    # Deduct points for warnings
    score -= len(warnings) * 5  # Each warning costs 5 points
    
    score = max(0, min(100, score))
    
    # Determine validity
    is_valid = len(issues) == 0 and score >= 70
    
    return {
        "valid": is_valid,
        "score": score,
        "file_type": file_type,
        "issues": issues,
        "warnings": warnings,
        "recommendations": recommendations,
        "summary": _generate_summary(issues, warnings, is_valid, score)
    }


def _generate_summary(issues: List[str], warnings: List[str], is_valid: bool, score: int) -> str:
    """Generate human-readable summary."""
    if is_valid and score >= 90:
        return "Resume is ATS-friendly and should parse correctly in most ATS systems."
    elif is_valid:
        return f"Resume is mostly ATS-friendly (score: {score}/100). Minor improvements recommended."
    elif score >= 50:
        return f"Resume has some ATS compatibility issues (score: {score}/100). Review recommendations."
    else:
        return f"Resume has significant ATS compatibility issues (score: {score}/100). Major improvements needed."

