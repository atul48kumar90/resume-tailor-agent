from docx import Document

def export_docx(resume: dict, path: str):
    doc = Document()

    if resume.get("summary"):
        doc.add_paragraph(resume["summary"])

    for exp in resume.get("experience", []):
        doc.add_heading(exp["title"], level=2)
        for b in exp.get("bullets", []):
            doc.add_paragraph(b, style="List Bullet")

    if resume.get("skills"):
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(resume["skills"]))

    doc.save(path)
