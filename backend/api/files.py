# api/files.py
from fastapi import UploadFile
import pdfplumber
from docx import Document
import re
import os
import tempfile
from typing import IO
from concurrent.futures import ThreadPoolExecutor
import asyncio
from core.security import (
    sanitize_filename,
    validate_file_content,
    validate_file_path,
)
from core.file_security import (
    validate_file_security,
    get_max_file_size_for_tier,
    scan_file,
)


# ======================================================
# File Type Detection (Magic Bytes)
# ======================================================

def _detect_file_type(file_obj: IO) -> str:
    """
    Detect file type by reading magic bytes (file signature).
    More secure than relying on file extension.
    
    Returns:
        File type: 'pdf', 'docx', 'txt', or 'unknown'
    """
    current_pos = file_obj.tell()
    file_obj.seek(0)
    
    try:
        # Read first few bytes to detect file type
        header = file_obj.read(8)
        file_obj.seek(current_pos)  # Reset position
        
        if not header:
            return 'unknown'
        
        # PDF: starts with %PDF
        if header.startswith(b'%PDF'):
            return 'pdf'
        
        # DOCX: is a ZIP file (starts with PK\x03\x04)
        # DOCX files are ZIP archives containing XML
        if header.startswith(b'PK\x03\x04'):
            # Check if it's actually a DOCX by looking for word/ in the ZIP
            # For now, we'll trust the PK header and validate when opening
            return 'docx'
        
        # TXT: try to decode as UTF-8 (heuristic)
        try:
            header.decode('utf-8')
            # If it's mostly printable ASCII/UTF-8, likely text
            if all(32 <= b <= 126 or b in [9, 10, 13] for b in header[:7]):
                return 'txt'
        except (UnicodeDecodeError, ValueError):
            pass
        
        return 'unknown'
    except Exception:
        file_obj.seek(current_pos)
        return 'unknown'


def _validate_file_type(file_obj: IO, expected_type: str, filename: str) -> bool:
    """
    Validate that file content matches expected type.
    
    Args:
        file_obj: File object
        expected_type: Expected file type ('pdf', 'docx', 'txt')
        filename: Original filename for error messages
    
    Returns:
        True if file type matches, False otherwise
    
    Raises:
        ValueError: If file type doesn't match
    """
    detected_type = _detect_file_type(file_obj)
    
    if detected_type == 'unknown':
        raise ValueError(
            f"Could not detect file type for {filename}. "
            f"File may be corrupted or in an unsupported format."
        )
    
    if detected_type != expected_type:
        raise ValueError(
            f"File type mismatch for {filename}. "
            f"Extension suggests {expected_type}, but file content indicates {detected_type}. "
            f"Please ensure the file is actually a {expected_type} file."
        )
    
    return True


# ======================================================
# Public API
# ======================================================

# Thread pool for blocking I/O operations
_file_executor = ThreadPoolExecutor(max_workers=4, thread_name_prefix="file_extract")


async def extract_text_async(file, max_size_bytes: int = None) -> str:
    """
    Async version of extract_text.
    Runs blocking file operations in thread pool.
    """
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(_file_executor, extract_text, file, max_size_bytes)


def extract_text(file, max_size_bytes: int = None) -> str:
    """
    Accepts:
    - FastAPI UploadFile
    - File opened via open(...)
    
    Args:
        file: File object to extract text from
        max_size_bytes: Maximum file size in bytes (None = no limit)
    
    Returns:
        Extracted and normalized text
    
    Raises:
        ValueError: If file type is unsupported or file is too large
    
    Uses caching to avoid re-extracting and re-normalizing the same file content.
    """
    import logging
    import hashlib
    from core.settings import MAX_FILE_SIZE_BYTES, CACHE_NORMALIZED_TTL
    from core.cache import (
        get_cached_extracted_text,
        set_cached_extracted_text,
    )
    
    logger = logging.getLogger(__name__)
    
    if hasattr(file, "filename"):
        # Sanitize filename to prevent path traversal
        try:
            filename = sanitize_filename(file.filename).lower()
        except ValueError as e:
            raise ValueError(f"Invalid filename: {str(e)}")
        file_obj = file.file
        # Check file size for UploadFile
        if hasattr(file, "size") and file.size:
            file_size = file.size
        else:
            # Read position to check size
            current_pos = file_obj.tell()
            file_obj.seek(0, 2)  # Seek to end
            file_size = file_obj.tell()
            file_obj.seek(current_pos)  # Reset position
    else:
        filename = os.path.basename(file.name).lower()
        file_obj = file
        # Check file size for regular file
        current_pos = file_obj.tell()
        file_obj.seek(0, 2)
        file_size = file_obj.tell()
        file_obj.seek(current_pos)

    # Determine user tier (default to 'free' if not provided)
    user_tier = getattr(file, 'user_tier', 'free') if hasattr(file, 'user_tier') else 'free'
    
    # Get tier-based file size limit
    max_size = max_size_bytes or get_max_file_size_for_tier(user_tier)
    if file_size > max_size:
        max_size_mb = max_size / (1024 * 1024)
        file_size_mb = file_size / (1024 * 1024)
        raise ValueError(
            f"File size ({file_size_mb:.2f} MB) exceeds maximum allowed size "
            f"({max_size_mb:.2f} MB) for {user_tier} tier. "
            f"Please upgrade your account or reduce file size."
        )
    
    # Determine expected file type
    expected_type = None
    if filename.endswith(".pdf"):
        expected_type = "pdf"
    elif filename.endswith(".docx"):
        expected_type = "docx"
    elif filename.endswith(".txt"):
        expected_type = "txt"
    
    # Comprehensive security validation (if expected_type is known)
    if expected_type:
        # Read file content for security checks
        file_obj.seek(0)
        file_content = file_obj.read(min(1024 * 1024, file_size))  # Read first 1MB for scanning
        file_obj.seek(0)  # Reset for extraction
        
        # Create temporary file for virus scanning if needed
        temp_file_path = None
        try:
            # For security validation, we need the full file
            # Create temp file if file is small enough
            if file_size < 10 * 1024 * 1024:  # Only for files < 10MB
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{expected_type}") as tmp:
                    file_obj.seek(0)
                    tmp.write(file_obj.read())
                    temp_file_path = tmp.name
                    file_obj.seek(0)  # Reset again
            
            # Perform comprehensive security check
            if temp_file_path:
                is_safe, error_msg, security_details = validate_file_security(
                    temp_file_path,
                    expected_type,
                    user_tier=user_tier,
                    file_content=file_content,
                    file_size=file_size
                )
                
                if not is_safe:
                    # Clean up temp file
                    try:
                        os.unlink(temp_file_path)
                    except:
                        pass
                    raise ValueError(error_msg or "File failed security validation")
                
                # Clean up temp file after successful validation
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
        except ValueError:
            # Re-raise validation errors
            if temp_file_path:
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
            raise
        except Exception as e:
            # Log but don't fail on security check errors (graceful degradation)
            logger.warning(f"Security check error (continuing anyway): {e}")
            if temp_file_path:
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
    
    # Validate file content (beyond magic bytes) - additional validation
    try:
        if expected_type == "pdf":
            validate_file_content(file_obj, "pdf", max_size)
        elif expected_type == "docx":
            validate_file_content(file_obj, "docx", max_size)
        elif expected_type == "txt":
            validate_file_content(file_obj, "txt", max_size)
    except ValueError as e:
        raise ValueError(f"File content validation failed: {str(e)}")

    # Generate file hash for caching (for files < 1MB)
    # Larger files skip file-level cache but still use text-level normalization cache
    file_hash = None
    if file_size < 1024 * 1024:  # Files < 1MB: cache by file hash
        file_obj.seek(0)
        file_content = file_obj.read()
        file_hash = hashlib.sha256(file_content).hexdigest()
        file_obj.seek(0)  # Reset for extraction
        
        # Check cache for extracted and normalized text
        cached_text = get_cached_extracted_text(file_hash)
        if cached_text:
            logger.info(f"Resume text cache hit for {filename} (hash: {file_hash[:8]}...)")
            return cached_text
        
        logger.info(f"Resume text cache miss for {filename}, extracting...")
    else:
        logger.info(f"Processing large file {filename} ({file_size / (1024*1024):.2f} MB), using text-level cache only")

    # Determine expected file type from extension
    try:
        if filename.endswith(".pdf"):
            expected_type = "pdf"
            _validate_file_type(file_obj, "pdf", filename)
            text = _extract_pdf(file_obj)
        elif filename.endswith(".docx"):
            expected_type = "docx"
            _validate_file_type(file_obj, "docx", filename)
            text = _extract_docx(file_obj)
        elif filename.endswith(".txt"):
            expected_type = "txt"
            # For text files, we're more lenient (no magic bytes validation)
            text = _extract_txt(file_obj)
        else:
            raise ValueError(
                f"Unsupported file type: {filename}. "
                f"Supported formats: .pdf, .docx, .txt. "
                f"Please convert your file to one of these formats."
            )
        
        # Validate extracted text
        if not text or len(text.strip()) < 10:
            logger.warning(f"Extracted text is very short or empty from {filename}")
            raise ValueError(
                f"Could not extract meaningful text from {filename}. "
                f"File may be image-based, corrupted, or password-protected."
            )
        
        # Log extraction success
        logger.info(f"Successfully extracted {len(text)} characters from {filename}")
        
        # 🔥 SINGLE SOURCE OF TRUTH - Use ATS normalization (lowercase for matching)
        normalized_text = normalize_resume_text_for_ats(text)
        
        # Cache the final normalized text (keyed by file hash) if we have the hash
        # Note: normalize_resume_text_for_ats() also caches normalized text internally
        if file_hash:
            set_cached_extracted_text(file_hash, normalized_text, ttl=CACHE_NORMALIZED_TTL)
        
        return normalized_text
    
    except ValueError:
        # Re-raise ValueError as-is (these are user-facing errors)
        raise
    except Exception as e:
        # Wrap other exceptions with context
        logger.error(f"Unexpected error extracting text from {filename}: {e}", exc_info=True)
        raise ValueError(
            f"Failed to extract text from {filename}: {str(e)}. "
            f"Please ensure the file is not corrupted and is in a supported format."
        )


# ======================================================
# PDF Extraction (ROBUST)
# ======================================================

def _extract_pdf(file_obj: IO) -> str:
    """
    Enhanced PDF extraction with:
    - Table extraction
    - Multi-column layout handling
    - Fallback strategies
    - Better error handling
    """
    text_chunks = []
    
    try:
        # Primary method: pdfplumber (better for complex layouts)
        with pdfplumber.open(file_obj) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text_parts = []
                
                # 1. Extract regular text with layout preservation
                page_text = page.extract_text(
                    x_tolerance=2,
                    y_tolerance=2,
                    layout=True,
                )
                if page_text:
                    page_text_parts.append(page_text)
                
                # 2. Extract tables (common in resumes for skills, experience)
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        if table:
                            # Convert table to readable text
                            table_text = _format_table_text(table)
                            if table_text:
                                page_text_parts.append(f"\n{table_text}\n")
                
                # 3. Fallback: Try without layout if no text found
                if not page_text_parts:
                    page_text = page.extract_text()
                    if page_text:
                        page_text_parts.append(page_text)
                
                if page_text_parts:
                    text_chunks.append("\n".join(page_text_parts))
        
        if text_chunks:
            return "\n\n".join(text_chunks)
    
    except Exception as e:
        # Fallback to pypdf if pdfplumber fails
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"pdfplumber extraction failed, trying pypdf fallback: {e}")
        
        try:
            from pypdf import PdfReader
            file_obj.seek(0)  # Reset file pointer
            reader = PdfReader(file_obj)
            text_chunks = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_chunks.append(text)
            return "\n\n".join(text_chunks)
        except Exception as e2:
            logger.error(f"Both PDF extraction methods failed: {e2}")
            raise ValueError(f"Failed to extract text from PDF: {str(e2)}")
    
    # If we get here, no text was extracted
    raise ValueError("No text could be extracted from PDF. File may be image-based or corrupted.")


def _format_table_text(table: list) -> str:
    """
    Convert a table (list of lists) to readable text format.
    Formats tables in a way that preserves information for ATS parsing.
    """
    if not table or not table[0]:
        return ""
    
    # Find max width for each column
    max_widths = []
    for row in table:
        for i, cell in enumerate(row):
            if cell:
                cell_text = str(cell).strip()
                if i >= len(max_widths):
                    max_widths.append(len(cell_text))
                else:
                    max_widths[i] = max(max_widths[i], len(cell_text))
    
    # Format rows
    formatted_rows = []
    for row in table:
        if not row:
            continue
        
        # Join cells with separator, handling None values
        cells = [str(cell).strip() if cell else "" for cell in row]
        row_text = " | ".join(cells)
        if row_text.strip():
            formatted_rows.append(row_text)
    
    return "\n".join(formatted_rows)


# ======================================================
# DOCX Extraction (ENHANCED)
# ======================================================

def _extract_docx(file_obj: IO) -> str:
    """
    Enhanced DOCX extraction with:
    - Table extraction
    - Better structure preservation
    - Header/footer handling
    - Error handling
    """
    text_parts = []
    
    try:
        doc = Document(file_obj)
        
        # 1. Extract paragraphs (main content)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        
        # 2. Extract tables (common in resumes for skills, experience)
        for table in doc.tables:
            table_text = _extract_docx_table(table)
            if table_text:
                text_parts.append(f"\n{table_text}\n")
        
        # 3. Extract headers and footers (may contain contact info)
        # Note: python-docx doesn't directly support headers/footers,
        # but we can try to get them from sections
        for section in doc.sections:
            # Headers
            if section.header:
                header_text = "\n".join(
                    para.text.strip() 
                    for para in section.header.paragraphs 
                    if para.text.strip()
                )
                if header_text:
                    text_parts.insert(0, header_text)  # Add at beginning
            
            # Footers
            if section.footer:
                footer_text = "\n".join(
                    para.text.strip() 
                    for para in section.footer.paragraphs 
                    if para.text.strip()
                )
                if footer_text:
                    text_parts.append(footer_text)
        
        if not text_parts:
            raise ValueError("No text found in DOCX file")
        
        return "\n".join(text_parts)
    
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def _extract_docx_table(table) -> str:
    """
    Extract text from a DOCX table in a readable format.
    """
    rows_text = []
    for row in table.rows:
        cells_text = []
        for cell in row.cells:
            # Get all text from cell (may have multiple paragraphs)
            cell_text = "\n".join(
                para.text.strip() 
                for para in cell.paragraphs 
                if para.text.strip()
            )
            if cell_text:
                cells_text.append(cell_text)
        
        if cells_text:
            rows_text.append(" | ".join(cells_text))
    
    return "\n".join(rows_text)


# ======================================================
# TXT Extraction
# ======================================================

def _extract_txt(file_obj: IO) -> str:
    return file_obj.read().decode("utf-8", errors="ignore")


# ======================================================
# Resume Normalization (MANDATORY)
# ======================================================

def normalize_resume_text(text: str, preserve_case: bool = False) -> str:
    """
    Repairs ATS-breaking artifacts introduced by PDF/DOCX extraction.
    This function MUST be applied before ATS scoring.
    
    Args:
        text: Raw text from PDF/DOCX extraction
        preserve_case: If True, preserves original case. If False, lowercases (for ATS matching)
    
    Returns:
        Normalized text ready for ATS scoring
    """
    if not text:
        return ""

    original_text = text
    if not preserve_case:
        text = text.lower()

    # --- Fix spaced letters (J a v a → java or Java)
    if preserve_case:
        # Preserve case but fix spacing
        text = re.sub(r'(?<!\w)([A-Za-z])\s+([A-Za-z])', lambda m: m.group(1) + m.group(2), text)
    else:
        text = re.sub(r'(?<!\w)([a-z])\s+([a-z])', r'\1\2', text)

    # --- Fix common ATS term fragmentation (case-insensitive)
    regex_replacements = {
        r'r\s*e\s*s\s*t': 'rest',
        r'a\s*p\s*i': 'api',
        r'h\s*t\s*t\s*p': 'http',
        r'j\s*s\s*o\s*n': 'json',
        r's\s*q\s*l': 'sql',
        r'k\s*u\s*b\s*e\s*r\s*n\s*e\s*t\s*e\s*s': 'kubernetes',
        r'd\s*o\s*c\s*k\s*e\s*r': 'docker',
    }

    for pattern, replacement in regex_replacements.items():
        if preserve_case:
            # Preserve original case of the matched text
            def replace_preserve_case(match):
                matched = match.group(0)
                if matched.isupper():
                    return replacement.upper()
                elif matched[0].isupper():
                    return replacement.capitalize()
                return replacement
            text = re.sub(pattern, replace_preserve_case, text, flags=re.I)
        else:
            text = re.sub(pattern, replacement, text, flags=re.I)

    # --- Normalize known resume patterns
    direct_replacements = {
        "ensuRESTability": "ensure rest stability",
        "http s": "http",
        "https": "http",
        "data structure & algorithm": "data structures algorithms",
        "data structure and algorithm": "data structures algorithms",
        "distributed system": "distributed systems",
        "highly available solutions": "high availability",
        "code review & optimization": "code review optimization",
    }

    for src, tgt in direct_replacements.items():
        if preserve_case:
            # Case-insensitive replacement but preserve surrounding case
            text = re.sub(re.escape(src), tgt, text, flags=re.I)
        else:
            text = text.replace(src.lower(), tgt)

    # --- Normalize bullets & separators
    # Preserve table separators (|) but normalize bullets
    text = text.replace("•", " ").replace("●", " ").replace("▪", " ")
    text = text.replace("○", " ").replace("■", " ").replace("□", " ")
    
    # Preserve table separators - don't replace | if it looks like a table
    # (tables typically have | with text on both sides)
    # Only replace standalone | or | at line boundaries
    text = re.sub(r'(?<!\w)\|(?!\w)', ' ', text)  # Replace | not surrounded by word chars
    text = re.sub(r'\|{2,}', ' ', text)  # Replace multiple ||
    
    # Normalize dashes but preserve ranges (e.g., "2019-2021")
    text = re.sub(r'(?<!\d)-(?!\d)', ' ', text)  # Replace - not between digits

    # --- Collapse whitespace (but preserve line breaks for structure)
    # First, normalize line breaks
    text = re.sub(r'\r\n', '\n', text)  # Windows line breaks
    text = re.sub(r'\r', '\n', text)  # Old Mac line breaks
    
    # Collapse multiple spaces but preserve single spaces and newlines
    text = re.sub(r'[ \t]+', ' ', text)  # Collapse spaces/tabs
    text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 consecutive newlines
    
    # Final cleanup
    text = text.strip()

    return text


def normalize_resume_text_for_ats(text: str) -> str:
    """
    Normalizes text specifically for ATS scoring (lowercase).
    Use this when you need case-insensitive matching.
    
    Uses caching to avoid redundant normalization of identical text.
    """
    from core.cache import (
        get_cached_normalized_text,
        set_cached_normalized_text,
    )
    from core.settings import CACHE_NORMALIZED_TTL
    
    if not text or not text.strip():
        return ""
    
    # Check cache first (cache all texts for better performance)
    cached = get_cached_normalized_text(text)
    if cached:
        return cached
    
    # Normalize the text
    normalized = normalize_resume_text(text, preserve_case=False)
    
    # Cache the result (cache all normalized texts)
    set_cached_normalized_text(text, normalized, ttl=CACHE_NORMALIZED_TTL)
    
    return normalized


def normalize_resume_text_preserve_case(text: str) -> str:
    """
    Normalizes text while preserving original case.
    Use this when you need to display or export the text.
    """
    return normalize_resume_text(text, preserve_case=True)
