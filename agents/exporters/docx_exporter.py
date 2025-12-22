from docx import Document


def export_docx(resume: dict) -> bytes:
    doc = Document()

    if resume.get("summary"):
        doc.add_heading("SUMMARY", level=1)
        doc.add_paragraph(resume["summary"])

    if resume.get("experience"):
        doc.add_heading("EXPERIENCE", level=1)
        for exp in resume["experience"]:
            doc.add_paragraph(exp.get("title", ""), style="List Bullet")
            for b in exp.get("bullets", []):
                doc.add_paragraph(b, style="List Continue")

    if resume.get("skills"):
        doc.add_heading("SKILLS", level=1)
        doc.add_paragraph(", ".join(resume["skills"]))

    buf = bytes()
    from io import BytesIO
    f = BytesIO()
    doc.save(f)
    return f.getvalue()
